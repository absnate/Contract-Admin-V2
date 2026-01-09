import os
import json
import asyncio
from docx import Document
from fpdf import FPDF
import fitz  # PyMuPDF
from emergentintegrations.llm.chat import LlmChat, UserMessage

# Configuration
EMERGENT_LLM_KEY = os.environ.get("EMERGENT_LLM_KEY")
if not EMERGENT_LLM_KEY:
    print("Error: EMERGENT_LLM_KEY not set")
    exit(1)

# --- 1. Helper: Create Dummy Files ---
def create_dummy_docx(filename):
    doc = Document()
    doc.add_heading('Subcontract Agreement', 0)
    doc.add_paragraph('This Agreement is made between General Contractor Inc. ("GC") and Associated Building Specialties ("Subcontractor").')
    doc.add_heading('1. Indemnification', level=1)
    doc.add_paragraph('Subcontractor shall indemnify GC against any and all claims, regardless of negligence.')
    doc.add_heading('2. Payment', level=1)
    doc.add_paragraph('Payment shall be made within 7 days after GC receives payment from Owner (Pay-if-Paid).')
    doc.save(filename)
    print(f"Created {filename}")

def create_dummy_pdf(filename):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt="Subcontract Agreement (PDF Version)", ln=1, align="C")
    pdf.cell(200, 10, txt="1. Indemnification: Broad form indemnity applies.", ln=1)
    pdf.cell(200, 10, txt="2. Retainage: 10% retainage shall be withheld.", ln=1)
    pdf.output(filename)
    print(f"Created {filename}")

# --- 2. Extraction Logic ---
def extract_text_from_docx(path):
    print(f"Extracting from DOCX: {path}")
    doc = Document(path)
    text = []
    for para in doc.paragraphs:
        text.append(para.text)
    return "\n".join(text)

def extract_text_from_pdf(path):
    print(f"Extracting from PDF: {path}")
    text = ""
    with fitz.open(path) as doc:
        for page in doc:
            text += page.get_text()
    return text

# --- 3. LLM Analysis Logic (Using emergentintegrations) ---
async def analyze_contract(text, task_type="INITIAL_CONTRACT_REVIEW"):
    print(f"\n--- Analyzing with Gemini via EmergentIntegrations ({task_type}) ---")
    
    prompt = f"""
    You are the ABS Contract Admin Agent.
    TASK_TYPE: {task_type}
    
    Analyze the following contract text:
    {text}
    
    Output the result in JSON format with a field 'markdown_report' containing the readable report and 'structured_data' containing key fields.
    """
    
    try:
        # Initialize LlmChat
        chat = LlmChat(
            api_key=EMERGENT_LLM_KEY,
            session_id="test_session_001", # Dummy session
            system_message="You are a helpful contract analysis assistant."
        )
        
        # Configure for OpenAI (as fallback to prove connection)
        chat.with_model("openai", "gpt-4o")
        
        # Send message
        user_msg = UserMessage(text=prompt)
        response_text = await chat.send_message(user_msg)
        
        return response_text
    except Exception as e:
        print(f"LlmChat Error: {e}")
        return None

# --- 4. Main Execution Flow ---
async def main():
    # Setup
    docx_file = "test_contract.docx"
    pdf_file = "test_contract.pdf"
    
    create_dummy_docx(docx_file)
    create_dummy_pdf(pdf_file)
    
    # Test DOCX Extraction
    docx_text = extract_text_from_docx(docx_file)
    print(f"DOCX Text Preview: {docx_text[:100]}...")
    
    # Test LLM Analysis
    result_json = await analyze_contract(docx_text)
    
    if result_json:
        print("\n--- Analysis Result ---")
        print(result_json[:500] + "...")
        
        try:
            # Clean up markdown code blocks if present
            cleaned_json = result_json.replace("```json", "").replace("```", "").strip()
            data = json.loads(cleaned_json)
            if "markdown_report" in data:
                print("\nSuccessfully parsed JSON and found markdown_report.")
            else:
                print("\nJSON parsed but markdown_report missing.")
        except json.JSONDecodeError:
            print("\nFailed to parse JSON response (might be raw markdown).")
    else:
        print("\nAnalysis Failed.")

if __name__ == "__main__":
    asyncio.run(main())
