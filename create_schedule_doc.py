from docx import Document

doc = Document()
doc.add_heading('Contract with Schedule', 0)
doc.add_heading('Project Schedule', 1)
doc.add_paragraph('Project Start: Jan 1, 2024')
doc.add_paragraph('Milestone 1: Feb 1, 2024')
doc.add_paragraph('Project End: Mar 1, 2024')
doc.save('contract_with_schedule.docx')
