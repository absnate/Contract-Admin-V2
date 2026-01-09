import io
import fitz  # PyMuPDF
from docx import Document
from bson import ObjectId
from fpdf import FPDF

def extract_text_from_pdf(file_stream):
    """Extracts text from a PDF file stream."""
    text = ""
    try:
        pdf_bytes = file_stream.getvalue()
        with fitz.open(stream=pdf_bytes, filetype="pdf") as doc:
            for page in doc:
                text += page.get_text()
    except Exception as e:
        print(f"PDF Extraction Error: {e}")
        text = f"[Error extracting PDF: {str(e)}]"
    return text

def extract_text_from_docx(file_stream):
    """Extracts text from a DOCX file stream."""
    text = []
    try:
        doc = Document(file_stream)
        for para in doc.paragraphs:
            text.append(para.text)
        for table in doc.tables:
             for row in table.rows:
                 row_text = " | ".join([cell.text for cell in row.cells])
                 text.append(row_text)
    except Exception as e:
        print(f"DOCX Extraction Error: {e}")
        return f"[Error extracting DOCX: {str(e)}]"
    return "\n".join(text)

def create_pdf_from_text(text: str, title: str = "Extracted Content") -> bytes:
    """Creates a simple PDF from text string and returns bytes."""
    if not text or not text.strip():
        return b""

    try:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", "B", 16)
        pdf.cell(0, 10, title, ln=True, align="C")
        pdf.ln(10)
        
        pdf.set_font("Arial", size=12)
        
        # Clean text for FPDF compatibility (latin-1)
        # Remove unsupported characters
        safe_text = text.encode('latin-1', 'replace').decode('latin-1')
        
        pdf.multi_cell(0, 10, safe_text)
        
        # Output as latin-1 encoded bytes string
        return pdf.output(dest='S').encode('latin-1')
    except Exception as e:
        print(f"PDF Generation Error: {e}")
        return b""

def serialize_doc(doc):
    """Recursively converts ObjectId and datetime to strings for JSON serialization."""
    if isinstance(doc, list):
        return [serialize_doc(item) for item in doc]
    if isinstance(doc, dict):
        new_doc = {}
        for k, v in doc.items():
            if isinstance(v, ObjectId):
                new_doc[k] = str(v)
            elif hasattr(v, 'isoformat'): # datetime
                new_doc[k] = v.isoformat()
            elif isinstance(v, dict):
                new_doc[k] = serialize_doc(v)
            elif isinstance(v, list):
                new_doc[k] = serialize_doc(v)
            else:
                new_doc[k] = v
        return new_doc
    return doc
